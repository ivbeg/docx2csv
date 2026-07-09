# -*- coding: utf-8 -*-
"""Shared pytest fixtures for docx2csv tests."""

import pytest
from docx import Document


def _create_table_with_data(doc, data):
    """Helper to add a table with data to a document."""
    rows = len(data)
    cols = max(len(row) for row in data) if data else 1
    table = doc.add_table(rows=rows, cols=cols)
    for r, row_data in enumerate(data):
        for c, text in enumerate(row_data):
            table.cell(r, c).text = text
    return table


@pytest.fixture
def sample_docx(tmp_path):
    """Create a .docx file with 3 tables of varying sizes."""
    doc = Document()

    # Table 1: 2x2 simple
    _create_table_with_data(doc, [
        ["A", "B"],
        ["C", "D"],
    ])

    # Table 2: 3x3 with Cyrillic text (non-ASCII)
    _create_table_with_data(doc, [
        ["Я00", "Я01", "Я02"],
        ["Я10", "Я11", "Я12"],
        ["Я20", "Я21", "Я22"],
    ])

    # Table 3: 1x1 single cell
    _create_table_with_data(doc, [
        ["single"],
    ])

    path = tmp_path / "test.docx"
    doc.save(str(path))
    return str(path)


@pytest.fixture
def merged_docx(tmp_path):
    """Create a .docx file with merged cells."""
    doc = Document()

    # Table with a horizontal merge (2 cells merged)
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "merged_h"
    table.cell(0, 1).text = ""
    table.cell(0, 2).text = "right"
    table.cell(1, 0).text = "a"
    table.cell(1, 1).text = "b"
    table.cell(1, 2).text = "c"

    # Merge cells (0,0) and (0,1)
    table.cell(0, 0).merge(table.cell(0, 1))

    path = tmp_path / "merged.docx"
    doc.save(str(path))
    return str(path)


@pytest.fixture
def empty_docx(tmp_path):
    """Create a .docx file with no tables."""
    doc = Document()
    doc.add_paragraph("This document has no tables.")
    path = tmp_path / "empty.docx"
    doc.save(str(path))
    return str(path)


@pytest.fixture
def multline_docx(tmp_path):
    """Create a .docx file with multiline cell text."""
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Line1\nLine2\nLine3"
    table.cell(0, 1).text = "Single line"

    path = tmp_path / "multiline.docx"
    doc.save(str(path))
    return str(path)
