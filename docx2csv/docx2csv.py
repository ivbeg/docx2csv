# -*- coding: utf8 -*-

import csv
import xlwt

from docx import Document
from docx.table import _Cell
from docx.oxml.simpletypes import ST_Merge


def extract_table_old(table, verbose=True):
    """Extracts table data from table object"""
    results = []
    n = 0
    for row in table.rows:
        n += 1
        if verbose and n % 100 == 0: print 'Processed %d rows' % n
        r = []
        for cell in row.cells:
            r.append(cell.text.replace(u'\n', u' ').encode('utf8'))
        results.append(r)
    return results


def extract_table(table, verbose=True):
    """Extracts table data from table object"""
    results = []
    n = 0
    for tr in table._tbl.tr_lst:
        r = []
        for tc in tr.tc_lst:
            for grid_span_idx in range(tc.grid_span):
                if tc.vMerge == ST_Merge.CONTINUE:
                    r.append(results[n - 1][len(r) - 1])
                elif grid_span_idx > 0:
                    r.append(r[-1])
                else:
                    cell = _Cell(tc, table)
                    r.append(cell.text.replace('\n', ' ').encode('utf8'))
        results.append(r)
        n += 1
        if verbose and n % 100 == 0:
            print 'Processed %d rows' % n
    return results


def extract_docx_table(filename):
    """Extracts table from .DOCX files"""
    tables = []
    document = Document(filename)
#    print dir(document)
#    print document.tables
    n = 0
    for table in document.tables:
        n += 1
#        print '## TABLE %d ##' % (n)
        tdata = extract_table(table)
        tables.append(tdata)
    return tables


def store_table(tabdata, filename, format='csv'):
    """Saves table data as csv file"""
    if format == 'csv':
        f = file(filename, 'w')
        w = csv.writer(f, delimiter=',')
        for row in tabdata:
            w.writerow(row)
    elif format == 'xls':
        workbook = xlwt.Workbook()
        sheet = workbook.add_sheet('0')
        rn = 0
        for row in tabdata:
            cn = 0
            for c in row:
                sheet.write(rn, cn, c.decode('utf8'))
                cn += 1
            rn += 1
        workbook.save(filename)


def extract(filename, format='csv', sizefilter=0, singlefile=False):
    tables = extract_docx_table(filename)
    name = filename.rsplit('.', 1)[0]
    format = format.lower()
    n = 0
    lfilter = int(sizefilter)
    if singlefile:
        workbook = xlwt.Workbook()
        for t in tables:
            if lfilter >= len(t):
                print 'Table length %d instead of %d. Skipped' % (len(t), lfilter)
                continue
            n += 1
            sheet = workbook.add_sheet(str(n))
            rn = 0
            for row in t:
                cn = 0
                for c in row:
                    sheet.write(rn, cn, c.decode('utf8'))
                    cn += 1
                rn += 1
        destname = name + '.%s' % (format)
        workbook.save(destname)
        print destname, 'saved'
    else:
        for t in tables:
            if lfilter >= len(t):
                print 'Table length %d instead of %d. Skipped' % (len(t), lfilter)
                continue
            n += 1
            destname = name + '_%d.%s' % (n, format)
            store_table(t, destname, format)
            print destname, 'saved'
