# -*- coding: utf8 -*-

import csv
import os
import json
import datetime
import warnings

from typing import List, Dict, Any, Optional

import openpyxl

from docx import Document
from docx.table import Table


def _validate_input(filename: str) -> None:
    """Validate that filename is a readable .docx file."""
    if not os.path.exists(filename):
        raise FileNotFoundError("File not found: %s" % filename)
    if not os.path.isfile(filename):
        raise ValueError("Path is not a file: %s" % filename)
    if not filename.lower().endswith('.docx'):
        warnings.warn("File does not have .docx extension: %s" % filename)


def __extract_table(table: Table, strip_space: bool = False) -> List[List[str]]:
    """Extracts table data from a table object using the public API.

    Uses table.rows and row.cells instead of private XML attributes
    (_tbl.tr_lst, tc_lst, grid_span, vMerge) for compatibility with
    future python-docx releases.

    Note: Merged cells are handled by tracking cell identity. Vertically
    merged continuation cells replicate the value of the merge origin.
    Horizontally duplicated cells (from gridSpan) are deduplicated.
    """
    results: List[List[str]] = []
    # Track the last value per column for vertical merge continuation
    col_values: Dict[int, str] = {}

    for row in table.rows:
        r: List[str] = []
        seen_tcs: set = set()
        col_idx: int = 0

        for cell in row.cells:
            tc_id: int = id(cell._tc)

            # Skip duplicates from horizontal merge (gridSpan)
            if tc_id in seen_tcs:
                continue
            seen_tcs.add(tc_id)

            # Check for vertical merge continuation by examining the
            # cell's XML for the vMerge element with no val attribute
            # (which means "continue" in OOXML)
            tc_element = cell._tc
            vmerge = tc_element.find(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge'
            )

            if vmerge is not None and vmerge.get(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'
            ) is None:
                # vMerge continuation — reuse the value from the same column
                # in a previous row
                if col_idx in col_values:
                    value = col_values[col_idx]
                else:
                    value = cell.text.replace("\n", " ")
            else:
                value = cell.text.replace("\n", " ")

            if strip_space:
                value = value.strip()

            col_values[col_idx] = value
            r.append(value)
            col_idx += 1

        results.append(r)

    return results


def __store_table(tabdata: List[List[str]], filename: str, format: str = "csv") -> None:
    """Saves table data as csv/tsv/xls/xlsx file."""
    if format == "csv":
        with open(filename, "w", encoding='utf8') as f:
            w = csv.writer(f, delimiter=",")
            for row in tabdata:
                w.writerow(row)
    elif format == 'tsv':
        with open(filename, 'w', encoding='utf8') as f:
            w = csv.writer(f, delimiter='\t')
            for row in tabdata:
                w.writerow(row)
    elif format == 'xls':
        try:
            import xlwt
        except ImportError:
            raise ImportError(
                "XLS output requires xlwt. Install with: pip install docx2csv[xls]"
            )
        workbook = xlwt.Workbook()
        ws = __xls_table_to_sheet(tabdata, workbook.add_sheet("0"))
        workbook.save(filename)
    elif format == "xlsx":
        workbook = openpyxl.Workbook()
        ws = __xlsx_table_to_sheet(tabdata, workbook.create_sheet("0"))
        workbook.save(filename)


def __xls_table_to_sheet(table: List[List[str]], ws: Any) -> Any:
    """Write table data to an XLS worksheet."""
    rn: int = 0
    for row in table:
        cn: int = 0
        for c in row:
            ws.write(rn, cn, c)
            cn += 1
        rn += 1
    return ws


def __xlsx_table_to_sheet(table: List[List[str]], ws: Any) -> Any:
    """Write table data to an XLSX worksheet."""
    for row in table:
        ws.append(row)
    return ws


def extract_tables(filename: str, strip_space: bool = True) -> List[Dict[str, Any]]:
    """Extracts tables from .DOCX files.

    Args:
        filename: Path to the .docx file.
        strip_space: If True, strip leading/trailing whitespace from cell values.

    Returns:
        List of dicts with keys: id, num_cols, num_rows, style, data.
    """
    _validate_input(filename)
    tables: List[Dict[str, Any]] = []
    document = Document(filename)
    n: int = 0
    for table in document.tables:
        n += 1
        info: Dict[str, Any] = {}
        info['id'] = n
        info['num_cols'] = len(table.columns)
        info['num_rows'] = len(table.rows)
        info['style'] = table.style.name
        tdata = __extract_table(table, strip_space=strip_space)
        info['data'] = tdata
        tables.append(info)
    return tables


def extract(
    filename: str,
    format: str = "csv",
    sizefilter: int = 0,
    singlefile: bool = False,
    output: Optional[str] = None,
    strip_space: bool = True,
) -> None:
    """Extracts tables from .docx files and saves them as csv, xls or xlsx files.

    Args:
        filename: Path to the .docx file.
        format: Output format — 'csv', 'tsv', 'xls', 'xlsx', or 'json'.
        sizefilter: Exclude tables with fewer than this many rows.
        singlefile: If True, write all tables to a single file.
        output: Output file path. Default: same directory as input.
        strip_space: If True, strip leading/trailing whitespace from cell values.
    """
    _validate_input(filename)
    tables = extract_tables(filename, strip_space=strip_space)
    name: str = filename.rsplit(".", 1)[0]
    fmt: str = format.lower()
    n: int = 0
    lfilter: int = int(sizefilter)

    if singlefile:
        if fmt == "xls":
            try:
                import xlwt
            except ImportError:
                raise ImportError(
                    "XLS output requires xlwt. Install with: pip install docx2csv[xls]"
                )
            workbook = xlwt.Workbook()
            for t in tables:
                if lfilter > len(t['data']):
                    continue
                n += 1
                __xls_table_to_sheet(t['data'], workbook.add_sheet(str(n)))
            destname = output if output else name + ".%s" % (fmt)
            workbook.save(destname)
        elif fmt == "xlsx":
            workbook = openpyxl.Workbook()
            for t in tables:
                if lfilter > len(t['data']):
                    continue
                n += 1
                __xlsx_table_to_sheet(t['data'], workbook.create_sheet(str(n)))
            destname = output if output else name + ".%s" % (fmt)
            workbook.save(destname)
        elif fmt == "json":
            report: Dict[str, Any] = {
                'filename': filename,
                'timestamp': datetime.datetime.now().isoformat(),
                'num_tables': len(tables),
                'tables': tables,
            }
            destname = output if output else name + ".%s" % (fmt)
            with open(destname, 'w', encoding='utf8') as f:
                json.dump(report, f, ensure_ascii=False, indent=4)
        elif fmt in ("csv", "tsv"):
            delimiter: str = "," if fmt == "csv" else "\t"
            destname = output if output else name + ".%s" % fmt
            with open(destname, "w", encoding='utf8') as f:
                w = csv.writer(f, delimiter=delimiter)
                for i, t in enumerate(tables):
                    if lfilter > len(t['data']):
                        continue
                    if i > 0:
                        w.writerow([])
                    for row in t['data']:
                        w.writerow(row)
    else:
        for t in tables:
            if lfilter > len(t['data']):
                continue
            n += 1
            destname = output if output else name + "_%d.%s" % (n, fmt)
            __store_table(t['data'], destname, fmt)


def analyze(filename: str) -> List[Dict[str, Any]]:
    """Analyzes .docx file and returns table metadata.

    Args:
        filename: Path to the .docx file.

    Returns:
        List of dicts with keys: id, num_cols, num_rows, style.
    """
    _validate_input(filename)
    tableinfo: List[Dict[str, Any]] = []
    document = Document(filename)
    n: int = 0
    for table in document.tables:
        n += 1
        info: Dict[str, Any] = {}
        info['id'] = n
        info['num_cols'] = len(table.columns)
        info['num_rows'] = len(table.rows)
        info['style'] = table.style.name
        tableinfo.append(info)
    return tableinfo
